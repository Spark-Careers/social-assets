"""Spark Careers weekly content-calendar DOCX generator.

Builds spark-w{NN}-content-calendar.docx matching the Week 22 baseline:
cover page + week-at-a-glance table + per-post layouts.
"""

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


DAY_NAME = {
    "mon": "Monday", "tue": "Tuesday", "wed": "Wednesday",
    "thu": "Thursday", "fri": "Friday",
}
AUDIENCE_LABEL = {"b2b": "B2B", "b2c": "B2C"}


def _iso_week_monday(year: int, week_num: int) -> date:
    return date.fromisocalendar(year, week_num, 1)


def write_calendar_docx(captions: list[dict], year: int, week_num: int,
                         output_path: Path) -> None:
    week_label = f"{year}-W{week_num:02d}"
    nn = f"{week_num:02d}"
    monday = _iso_week_monday(year, week_num)
    friday = monday + timedelta(days=4)

    doc = Document()

    # --- Cover ---
    title = doc.add_heading("Spark Careers — Weekly Content Calendar", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Week {nn} · {monday.strftime('%b %d')} – {friday.strftime('%b %d, %Y')}")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x49, 0x9A, 0xA9)

    doc.add_paragraph("Two products, two audiences. RisePoint Careers brand on "
                       "Mission Monday; Spark Careers brand the rest of the week. "
                       "24 placements across LinkedIn (9), Facebook (10), Instagram (5).").runs[0].font.size = Pt(11)

    # --- Week at a glance ---
    doc.add_heading("Week at a glance", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Day", "Time", "Audience", "Theme", "Hook"]):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    for c in captions:
        d = _iso_week_monday(year, week_num) + timedelta(days=["mon", "tue", "wed", "thu", "fri"].index(c["day"]))
        time_str = "08:00" if c["audience"] == "b2b" else "12:00"
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(f"{DAY_NAME[c['day']]} {d.strftime('%b %d')}").font.size = Pt(10)
        row[1].paragraphs[0].add_run(time_str).font.size = Pt(10)
        row[2].paragraphs[0].add_run(AUDIENCE_LABEL[c["audience"]]).font.size = Pt(10)
        row[3].paragraphs[0].add_run(c["theme_tag"]).font.size = Pt(10)
        row[4].paragraphs[0].add_run(c["headline"]).font.size = Pt(10)

    doc.add_page_break()

    # --- Per-post layouts ---
    doc.add_heading("Posts in detail", level=1)
    for i, c in enumerate(captions, start=1):
        d = _iso_week_monday(year, week_num) + timedelta(days=["mon", "tue", "wed", "thu", "fri"].index(c["day"]))
        time_str = "08:00" if c["audience"] == "b2b" else "12:00"
        utm = f"?utm_source=social&utm_medium=post&utm_campaign=w{nn}-{c['day']}-{c['audience']}"

        h2 = doc.add_heading(f"{i}. {c['theme_tag']} — {AUDIENCE_LABEL[c['audience']]}", level=2)

        meta = doc.add_paragraph()
        meta.add_run(f"Date: ").bold = True
        meta.add_run(f"{DAY_NAME[c['day']]} {d.strftime('%B %d, %Y')}    ")
        meta.add_run(f"Time: ").bold = True
        meta.add_run(f"{time_str} MT (America/Edmonton)    ")
        meta.add_run(f"Channels: ").bold = True
        meta.add_run(", ".join(ch.title() for ch in c.get("channels", [])))

        media = doc.add_paragraph()
        media.add_run("Media file: ").bold = True
        media.add_run(f"{week_label}-{c['day']}-{c['audience']}.png")

        url_p = doc.add_paragraph()
        url_p.add_run("Tracked URL: ").bold = True
        url_p.add_run(f"https://{c['url']}/{utm}")

        doc.add_paragraph()
        doc.add_paragraph("Headline").runs[0].bold = True
        doc.add_paragraph(c["headline"]).paragraph_format.left_indent = Cm(0.5)

        if c.get("subline"):
            doc.add_paragraph("Subline").runs[0].bold = True
            doc.add_paragraph(c["subline"]).paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph("LinkedIn caption").runs[0].bold = True
        if "linkedin" in c.get("channels", []):
            p = doc.add_paragraph(c.get("caption_linkedin", ""))
            p.paragraph_format.left_indent = Cm(0.5)
        else:
            p = doc.add_paragraph("(skipped on LinkedIn)")
            p.paragraph_format.left_indent = Cm(0.5)
            p.runs[0].italic = True

        doc.add_paragraph("Facebook caption").runs[0].bold = True
        p = doc.add_paragraph(c.get("caption_facebook", ""))
        p.paragraph_format.left_indent = Cm(0.5)

        if c["audience"] == "b2c" and "instagram" in c.get("channels", []):
            doc.add_paragraph("Instagram caption").runs[0].bold = True
            p = doc.add_paragraph(c.get("caption_instagram", ""))
            p.paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph("Alt text").runs[0].bold = True
        alt = (f"{c['theme_tag']} graphic for {AUDIENCE_LABEL[c['audience']]} audience. "
               f"Headline: {c['headline']} 1080x1350.")
        p = doc.add_paragraph(alt)
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].italic = True

        doc.add_paragraph()  # spacer

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    import argparse
    import json

    parser = argparse.ArgumentParser()
    parser.add_argument("--captions", type=Path, required=True)
    parser.add_argument("--year", type=int, required=True)
    parser.add_argument("--week", type=int, required=True)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    captions = json.loads(args.captions.read_text(encoding="utf-8"))
    write_calendar_docx(captions, args.year, args.week, args.out)
    print(f"Wrote {args.out}")
